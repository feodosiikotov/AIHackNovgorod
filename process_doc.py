from catboost import CatBoostClassifier
import re
from docx import Document
import pandas as pd
import pickle
import nltk
from pymystem3 import Mystem
from string import punctuation
from docx.text.paragraph import Paragraph
from pymorphy2 import MorphAnalyzer
from nltk.corpus import stopwords
import re

factors_dict = {0: 'широта дискреционных полномочий\n', 
        1: 'определение компетенции по формуле "вправе"\n', 
        2: 'выборочное изменение объема прав\n',
        3: 'чрезмерная свобода подзаконного нормотворчества\n', 
        4: 'принятие нормативного правового акта за пределами компетенции\n', 
        5: 'заполнение законодательных пробелов при помощи подзаконных актов в отсутствие законодательной делегации соответствующих полномочий\n',
        6: 'отсутствие или неполнота административных процедур\n', 
        7: 'отказ от конкурсных (аукционных) процедур\n', 
        8: 'нормативные коллизии - противоречия, в том числе внутренние, между нормами\n',
        9: 'наличие завышенных требований к лицу, предъявляемых для реализации принадлежащего ему права\n', 
        10: 'злоупотребление правом заявителя государственными органами, органами местного самоуправления или организациями (их должностными лицами)\n',
        11: 'юридико-лингвистическая неопределенность\n'}

bin_m = CatBoostClassifier()
bin_m.load_model('bin_model.cbm')
color_m = CatBoostClassifier()
color_m.load_model('color_model.cbm')
multi_m = CatBoostClassifier()
multi_m.load_model('multi_model.cbm')

with open('svd.p', 'rb') as f:
    svd_bin = pickle.load(f)
with open('vectorizer.p', 'rb') as f:
    vectorizer_bin = pickle.load(f)
with open('tfidf_color.pickle', 'rb') as f:
    vectorizer_color = pickle.load(f)

def preprocess(text):
  morph_analyzer = MorphAnalyzer()
  stop_words = stopwords.words("russian")
  text = text.lower()
  text = re.sub('\s+', ' ', re.sub(r'[^А-Яа-яёЁ]', ' ', text).strip())
  text = [morph_analyzer.parse(word)[0].normal_form for word in text.split()]
  text = [word for word in text if word not in stop_words]
  return ' '.join(text)


def process_doc(filename, save_filename):
    print(0)
    print(filename)
    doc = Document(filename)
    print(1)
    has_factor_in_doc = False
    factors = []
    for i, paragraph in enumerate(doc.paragraphs):
        X = preprocess(paragraph.text)
        X = vectorizer_bin.transform([X])
        X = svd_bin.transform(X)
        has_factor = bin_m.predict(X)[0].astype(bool)
        if has_factor:
            has_factor_in_doc = True
            factor = multi_m.predict(X)
            factors += [factor[0][0]]
            print(factor[0][0])
            for j, run in enumerate(paragraph.runs):
                X = preprocess(run.text)
                X = vectorizer_color.transform([X])
                has_color = color_m.predict(X)[0].astype(bool)
                if has_color:
                    doc.paragraphs[i].runs[j].font.highlight_color = 6
                else:
                    doc.paragraphs[i].runs[j].font.highlight_color = 7
    
    doc.save(save_filename)
    if not has_factor_in_doc:
        'Коррупциогенные факторы не были выявлены'
    else:
        return_str = 'Возможно, данный файл содержит следующие корупциогенные факторы:\n'
        factors = list(set(factors))
        for factor in factors:
            return_str += factors_dict[factor]
    return return_str

# process_doc('Edition_text.docx', 'new_doc.docx')



